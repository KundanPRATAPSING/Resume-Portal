from datetime import datetime
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "client" / "public" / "docs" / "ResumePortal_Interview_Documentation.md"
TARGET_DOCX = ROOT / "client" / "public" / "docs" / "ResumePortal_Interview_Documentation.docx"


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(15, 43, 91)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(25, 70, 135)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(35, 98, 171)


def add_cover_page(doc: Document) -> None:
    title = doc.add_paragraph("Resume Portal")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(15, 43, 91)
    set_paragraph_spacing(title, before=12, after=10, line=1.0)

    subtitle = doc.add_paragraph("Complete Interview Documentation")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(25, 70, 135)
    set_paragraph_spacing(subtitle, before=0, after=24, line=1.0)

    meta = doc.add_paragraph(
        "Feature-by-feature implementation, architecture logic, code mapping, and interview-ready explanations."
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(meta, before=0, after=18, line=1.15)

    generated = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%d %b %Y, %I:%M %p')}")
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated.runs[0].italic = True
    set_paragraph_spacing(generated, before=0, after=0, line=1.0)

    doc.add_page_break()


def add_inline_markdown(paragraph, text: str) -> None:
    # Supports inline code (`code`) and bold (**text**) for readable DOCX output.
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(10)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def parse_md_line(doc: Document, line: str) -> None:
    clean = line.rstrip("\n").strip()

    if not clean:
        p = doc.add_paragraph("")
        set_paragraph_spacing(p, before=0, after=4, line=1.0)
        return

    if clean == "---":
        p = doc.add_paragraph(" ")
        set_paragraph_spacing(p, before=3, after=6, line=1.0)
        return

    if clean.startswith("# "):
        p = doc.add_heading(clean[2:].strip(), level=1)
        set_paragraph_spacing(p, before=10, after=8, line=1.0)
        return

    if clean.startswith("## "):
        p = doc.add_heading(clean[3:].strip(), level=2)
        set_paragraph_spacing(p, before=8, after=5, line=1.0)
        return

    if clean.startswith("### "):
        p = doc.add_heading(clean[4:].strip(), level=3)
        set_paragraph_spacing(p, before=6, after=4, line=1.0)
        return

    if re.match(r"^\d+\.\s+", clean):
        text = re.sub(r"^\d+\.\s+", "", clean)
        p = doc.add_paragraph(style="List Number")
        add_inline_markdown(p, text)
        set_paragraph_spacing(p, before=0, after=4, line=1.1)
        return

    if clean.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        add_inline_markdown(p, clean[2:].strip())
        set_paragraph_spacing(p, before=0, after=4, line=1.1)
        return

    p = doc.add_paragraph()
    add_inline_markdown(p, clean)
    set_paragraph_spacing(p, before=0, after=8, line=1.15)


def main() -> None:
    if not SOURCE_MD.exists():
        raise FileNotFoundError(f"Missing source markdown: {SOURCE_MD}")

    doc = Document()
    style_document(doc)
    add_cover_page(doc)

    for raw in SOURCE_MD.read_text(encoding="utf-8").splitlines():
        parse_md_line(doc, raw)

    doc.save(str(TARGET_DOCX))
    print(f"Generated: {TARGET_DOCX}")


if __name__ == "__main__":
    main()
